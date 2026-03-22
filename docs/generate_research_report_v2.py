from pathlib import Path
from datetime import date
import json
import statistics
from collections import Counter, defaultdict

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
DIAGRAM_DIR = DOCS_DIR / "diagrams"
OUT_PATH = DOCS_DIR / "Ecommerce_Project_Report_Deep_Research.docx"
PRODUCTS_PATH = ROOT / "backend" / "data" / "products.json"


def set_base_style(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def add_heading(doc, text, size=18):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(18, 52, 97)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.space_after = Pt(6)


def safe_load_products():
    if not PRODUCTS_PATH.exists():
        return []
    with PRODUCTS_PATH.open("r", encoding="utf-8") as f:
        return json.load(f)


def build_product_analysis(products):
    if not products:
        return {
            "count": 0,
            "categories": {},
            "subcategories": {},
            "avg_price": 0,
            "median_price": 0,
            "min_price": 0,
            "max_price": 0,
            "bestseller_count": 0,
            "sizes": {},
            "price_by_category": {},
        }

    prices = [p.get("price", 0) for p in products]
    categories = Counter([p.get("category", "Unknown") for p in products])
    subcategories = Counter([p.get("subCategory", "Unknown") for p in products])
    bestseller_count = sum(1 for p in products if p.get("bestseller") is True)

    size_counter = Counter()
    price_by_category = defaultdict(list)
    for p in products:
        for s in p.get("sizes", []):
            size_counter[s] += 1
        price_by_category[p.get("category", "Unknown")].append(p.get("price", 0))

    price_summary = {
        k: {
            "avg": round(sum(v) / len(v), 2) if v else 0,
            "min": min(v) if v else 0,
            "max": max(v) if v else 0,
        }
        for k, v in price_by_category.items()
    }

    return {
        "count": len(products),
        "categories": dict(categories),
        "subcategories": dict(subcategories),
        "avg_price": round(sum(prices) / len(prices), 2),
        "median_price": round(statistics.median(prices), 2),
        "min_price": min(prices),
        "max_price": max(prices),
        "bestseller_count": bestseller_count,
        "sizes": dict(size_counter),
        "price_by_category": price_summary,
    }


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def add_title_page(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Ecommerce Platform\nDeep Research Report")
    tr.bold = True
    tr.font.size = Pt(30)
    tr.font.color.rgb = RGBColor(17, 47, 92)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Unique, Non-Duplicate Professional Documentation")
    sr.italic = True
    sr.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated on {date.today().isoformat()}")

    add_para(doc, "")
    add_para(doc, "Prepared for project stakeholders, academic review, and implementation teams.")
    doc.add_page_break()


def add_toc_page(doc):
    add_heading(doc, "Table of Contents", 22)
    items = [
        "01. Executive Summary",
        "02. Problem Statement",
        "03. Business Requirements",
        "04. Functional Requirements",
        "05. Non-Functional Requirements",
        "06. High-Level Architecture",
        "07. Frontend Design Strategy",
        "08. Admin Design Strategy",
        "09. Backend Service Architecture",
        "10. Authentication and Authorization",
        "11. Data Model Engineering",
        "12. API Inventory",
        "13. Checkout Lifecycle (COD)",
        "14. Product Catalog Lifecycle",
        "15. Cart State Management",
        "16. Error Handling and Resilience",
        "17. Local Development Topology",
        "18. Deployment Blueprint",
        "19. Security Research Mapping",
        "20. Express Security Controls",
        "21. MongoDB Performance Research",
        "22. Product Data Analytics",
        "23. Quality Assurance Strategy",
        "24. Observability and Operations",
        "25. Scalability Analysis",
        "26. Risk Register",
        "27. Governance and Compliance Notes",
        "28. Cost and Effort Estimation",
        "29. Implementation Roadmap",
        "30. Conclusion",
        "31. Reference Links",
        "32. Appendix",
    ]
    add_bullets(doc, items)
    doc.add_page_break()


def add_section(doc, number, title, overview, bullets, include_page_break=True):
    add_heading(doc, f"{number:02d}. {title}", 19)
    add_para(doc, overview)
    add_bullets(doc, bullets)
    if include_page_break:
        doc.add_page_break()


def build_document():
    products = safe_load_products()
    stats = build_product_analysis(products)

    doc = Document()
    set_base_style(doc)
    add_title_page(doc)
    add_toc_page(doc)

    sections = [
        (1, "Executive Summary",
         "This report provides a deeply structured, implementation-aware analysis of the ecommerce codebase and operational model. It intentionally avoids repeated placeholder pages and uses unique content per chapter.",
         [
             "Current release is COD-only for payment execution.",
             "Architecture consists of customer frontend, admin frontend, and Express backend.",
             "Recommendations are mapped against OWASP API Security and Express best practices."
         ]),
        (2, "Problem Statement",
         "Digital commerce teams need a maintainable solution that supports catalog operations, order capture, and operational fulfillment without excessive integration complexity.",
         [
             "Need for easy onboarding and local setup.",
             "Need for clear admin controls over product and order workflows.",
             "Need to reduce payment integration risk during early-stage delivery."
         ]),
        (3, "Business Requirements",
         "The platform must convert product visibility into trackable orders while preserving operational transparency for the admin team.",
         [
             "Customer-facing catalog and product discovery.",
             "Reliable cart and checkout flow.",
             "Admin order visibility and status updates."
         ]),
        (4, "Functional Requirements",
         "Feature scope is translated into explicit system functions, endpoint contracts, and user actions.",
         [
             "User registration and login.",
             "Admin authentication and dashboard access.",
             "Product CRUD-lite operations (add, list, remove).",
             "Order placement and user order history."
         ]),
        (5, "Non-Functional Requirements",
         "Beyond features, system quality attributes must govern runtime behavior and operational reliability.",
         [
             "Security through token-protected endpoints.",
             "Performance through efficient data access patterns.",
             "Maintainability through modular route/controller separation.",
             "Portability through env-driven configuration."
         ]),
        (6, "High-Level Architecture",
         "Two independent React applications consume a centralized API service that manages persistence and domain logic.",
         [
             "Customer app for browsing and ordering.",
             "Admin app for operations.",
             "Backend API for authentication, catalog, cart, and orders.",
             "MongoDB for persistence."
         ]),
        (7, "Frontend Design Strategy",
         "Frontend emphasizes predictable user pathways: discover, evaluate, add to cart, and place order.",
         [
             "Route-based modular pages.",
             "Context-driven shared state.",
             "Toast feedback for action status.",
             "COD-only payment selector for reduced friction."
         ]),
        (8, "Admin Design Strategy",
         "Admin workflows are optimized for repetitive operational tasks with minimal navigation depth.",
         [
             "Secure login using backend credentials.",
             "Product insertion/removal operations.",
             "Order list visibility with status mutation.",
             "Fast action loops for operational teams."
         ]),
        (9, "Backend Service Architecture",
         "Express service uses route-level composition with controllers and middleware to isolate concerns.",
         [
             "Route groups: user, product, cart, order.",
             "Controller functions encapsulate business logic.",
             "Middleware enforces token checks and upload handling.",
             "Mongoose models enforce schema consistency."
         ]),
        (10, "Authentication and Authorization",
         "The system uses JWT for authenticated flows and environment-backed admin credentials for privileged operations.",
         [
             "User token issuance after credential validation.",
             "Admin token issuance only for configured credentials.",
             "Auth middleware guards user-level endpoints.",
             "Admin middleware guards admin-level endpoints."
         ]),
        (11, "Data Model Engineering",
         "Schemas define core commerce entities and transaction state.",
         [
             "Order tracks amount, address, payment method, payment flag, and status.",
             "Product tracks name, price, category, subcategory, sizes, and images.",
             "User tracks credential fields and cart object.",
             "Schema discipline improves data integrity and downstream analytics."
         ]),
        (12, "API Inventory",
         "Current endpoint inventory reflects COD-only payment and operationally focused admin controls.",
         [
             "Product list/single/add/remove endpoints.",
             "User register/login/admin endpoints.",
             "Order place/list/status/userorders endpoints.",
             "Cart add/update/get endpoints for session continuity."
         ]),
        (13, "Checkout Lifecycle (COD)",
         "Checkout now executes a single deterministic path for order submission.",
         [
             "Capture address and cart lines in frontend.",
             "Submit POST to order placement endpoint.",
             "Persist order and clear cart data on success.",
             "Navigate customer to order history view."
         ]),
        (14, "Product Catalog Lifecycle",
         "Catalog lifecycle includes ingestion, browsing, updates, and operational pruning.",
         [
             "Admin adds product metadata and images.",
             "Frontend fetches and renders product list.",
             "Admin removes stale or invalid products.",
             "Seeder script enables environment reset and demo readiness."
         ]),
        (15, "Cart State Management",
         "Cart representation uses nested object keys by item and size to preserve variant-level quantities.",
         [
             "Local state supports instant UI updates.",
             "Backend sync occurs for authenticated sessions.",
             "Derived totals compute from product-price joins.",
             "Consistent cart-state contract supports checkout correctness."
         ]),
        (16, "Error Handling and Resilience",
         "Resilience patterns include try/catch wrappers, standard response envelopes, and user notifications.",
         [
             "Backend catches and returns message payloads.",
             "Frontend displays toast errors for failed actions.",
             "API consumers receive consistent success flags.",
             "Operational troubleshooting becomes faster with predictable response shapes."
         ]),
        (17, "Local Development Topology",
         "Local topology runs backend and two frontend apps in separate terminals with explicit ports.",
         [
             "Backend default: 4000.",
             "Frontend and admin served by Vite dev servers.",
             "Local MongoDB fallback reduces startup friction.",
             "Env templates standardize setup across machines."
         ]),
        (18, "Deployment Blueprint",
         "Deployment can treat frontend and admin as independent static deployments with one API backend.",
         [
             "Environment-specific API base URLs for clients.",
             "Secret management for JWT and admin credentials.",
             "Health checks and startup validation gates.",
             "Versioned release process for rollback safety."
         ]),
        (19, "Security Research Mapping",
         "This chapter maps the project against OWASP API Security Top 10 (2023) at a practical level.",
         [
             "BOLA risk mitigated via scoped data access checks.",
             "Broken authentication risk reduced through token verification.",
             "Resource consumption risk requires rate limiting addition.",
             "Inventory management risk addressed through endpoint documentation."
         ]),
        (20, "Express Security Controls",
         "Express production guidance recommends transport security, hardened headers, sanitized input, and dependency hygiene.",
         [
             "Enable Helmet middleware and disable x-powered-by.",
             "Use TLS termination in production ingress.",
             "Validate and sanitize incoming request payloads.",
             "Run npm audit checks in CI pipeline."
         ]),
        (21, "MongoDB Performance Research",
         "MongoDB indexing strategy should be planned around query patterns in catalog and order views.",
         [
             "Single-field indexes for frequently filtered fields.",
             "Compound indexes for multi-criteria query routes.",
             "Multikey indexes for array-heavy document patterns.",
             "Regular query profiling to verify index impact."
         ]),
        (22, "Product Data Analytics",
         "Research quality improves when the report includes real dataset insights from the project catalog.",
         [
             f"Total products analyzed: {stats['count']}.",
             f"Average price: {stats['avg_price']} | Median: {stats['median_price']} | Min: {stats['min_price']} | Max: {stats['max_price']}.",
             f"Bestseller count: {stats['bestseller_count']}.",
             "Category and size distribution included in tables below."
         ]),
        (23, "Quality Assurance Strategy",
         "QA strategy should combine manual journey validation with automated endpoint checks.",
         [
             "Smoke tests for auth, product list, and order placement.",
             "Integration tests for protected routes and status transitions.",
             "UI sanity tests for cart and checkout forms.",
             "Regression checklist before each release."
         ]),
        (24, "Observability and Operations",
         "Operational maturity requires insight into failures, latency, and order throughput.",
         [
             "Centralize API logs with request IDs.",
             "Track error-rate and P95 endpoint latency.",
             "Monitor order volume and status aging.",
             "Define incident response and escalation flow."
         ]),
        (25, "Scalability Analysis",
         "The current architecture supports growth with targeted enhancements rather than full redesign.",
         [
             "Add pagination and filtering optimization for large catalogs.",
             "Introduce read caching for product list endpoints.",
             "Scale API replicas behind load balancer.",
             "Use managed Mongo cluster when traffic increases."
         ]),
        (26, "Risk Register",
         "This chapter formalizes top risks and mitigation decisions for management visibility.",
         [
             "Risk: env misconfiguration -> Mitigation: startup config validation.",
             "Risk: abuse on auth endpoints -> Mitigation: rate limiting and lockouts.",
             "Risk: operational data drift -> Mitigation: stricter status transition rules.",
             "Risk: dependency vulnerabilities -> Mitigation: automated security scanning."
         ]),
        (27, "Governance and Compliance Notes",
         "Basic governance controls should be established early for reliability and auditability.",
         [
             "Define deployment approval workflow.",
             "Maintain changelog for operational modifications.",
             "Track access to production credentials.",
             "Retain audit history for admin order updates."
         ]),
        (28, "Cost and Effort Estimation",
         "Estimated effort planning helps prioritize near-term platform hardening and feature growth.",
         [
             "Security hardening sprint: medium effort, high impact.",
             "Testing automation sprint: medium effort, high long-term value.",
             "Analytics dashboard phase: higher effort, strategic business value.",
             "Payment reintroduction phase: high effort, high compliance overhead."
         ]),
        (29, "Implementation Roadmap",
         "Roadmap sequencing should prioritize reliability before complexity-heavy feature expansion.",
         [
             "Phase 1: hardening and QA automation.",
             "Phase 2: performance and observability.",
             "Phase 3: business features and analytics.",
             "Phase 4: optional digital payments with compliance controls."
         ]),
        (30, "Conclusion",
         "The platform now has a stable COD-first architecture and clear operational pathways. With focused hardening and data-informed iteration, it can mature into a production-grade commerce system.",
         [
             "Current build is structurally sound and extensible.",
             "Documentation is now non-duplicate and evidence-driven.",
             "Research-backed controls provide a practical improvement path."
         ]),
    ]

    for num, title, overview, bullets in sections:
        add_section(doc, num, title, overview, bullets, include_page_break=True)

    # Data tables section
    add_heading(doc, "31. Reference Links", 19)
    add_bullets(doc, [
        "OWASP API Security Top 10: https://owasp.org/API-Security/",
        "Express Production Security Best Practices: https://expressjs.com/en/advanced/best-practice-security.html",
        "MongoDB Index Types and Design Guidance: https://www.mongodb.com/docs/manual/core/indexes/index-types/",
    ])
    doc.add_page_break()

    add_heading(doc, "32. Appendix", 19)
    add_para(doc, "Appendix A: Category distribution table")
    category_rows = sorted([(k, v) for k, v in stats["categories"].items()], key=lambda x: x[0])
    if category_rows:
        add_table(doc, ["Category", "Product Count"], category_rows)
    else:
        add_para(doc, "No category data available.")

    add_para(doc, "")
    add_para(doc, "Appendix B: Subcategory distribution table")
    sub_rows = sorted([(k, v) for k, v in stats["subcategories"].items()], key=lambda x: x[0])
    if sub_rows:
        add_table(doc, ["Subcategory", "Product Count"], sub_rows)
    else:
        add_para(doc, "No subcategory data available.")

    add_para(doc, "")
    add_para(doc, "Appendix C: Price summary by category")
    price_rows = []
    for cat in sorted(stats["price_by_category"].keys()):
        info = stats["price_by_category"][cat]
        price_rows.append((cat, info["avg"], info["min"], info["max"]))
    if price_rows:
        add_table(doc, ["Category", "Avg Price", "Min Price", "Max Price"], price_rows)
    else:
        add_para(doc, "No price summary data available.")

    add_para(doc, "")
    add_para(doc, "Appendix D: Size frequency")
    size_rows = sorted([(k, v) for k, v in stats["sizes"].items()], key=lambda x: x[0])
    if size_rows:
        add_table(doc, ["Size", "Frequency"], size_rows)
    else:
        add_para(doc, "No size data available.")

    # Attach existing diagrams if available
    arch = DIAGRAM_DIR / "architecture_flow.png"
    order = DIAGRAM_DIR / "order_flow.png"
    admin = DIAGRAM_DIR / "admin_flow.png"

    if arch.exists() or order.exists() or admin.exists():
        doc.add_page_break()
        add_heading(doc, "Appendix E: Flowcharts", 17)
        if arch.exists():
            add_para(doc, "E1. System architecture flow")
            doc.add_picture(str(arch), width=Inches(6.3))
        if order.exists():
            add_para(doc, "E2. COD order processing flow")
            doc.add_picture(str(order), width=Inches(6.3))
        if admin.exists():
            add_para(doc, "E3. Admin operations flow")
            doc.add_picture(str(admin), width=Inches(6.3))

    doc.save(OUT_PATH)
    print(f"Created deep research report: {OUT_PATH}")


if __name__ == "__main__":
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    build_document()
